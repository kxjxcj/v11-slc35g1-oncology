#!/usr/bin/env python3
"""
Polish v8 manuscript:
1. Cleanup version markers (v11/v6.7/v8 → empty or appropriate)
2. Remove AI markers (★ / 喵喵喵 / Mavis / 老刀 trigger words / AI emoji)
3. Move figures inline to corresponding Results subsections (§4.1-§4.5)
4. Strip [E3 Visium] / [E6 MR] / [E8 PDO] / [v6.7 NEW] tags from captions

Output: ~/Desktop/v11_final/v11_final_v8_manuscript.docx (overwrites)
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from copy import deepcopy
from pathlib import Path
import re

V8 = '/Users/chen/Desktop/v11_final/v11_final_v8_manuscript.docx'

print(f'[v8 polish] loading {V8}')
doc = Document(V8)

# ============ STEP 1: TEXT CLEANUP ============
# Patterns to remove / replace
cleanup_patterns = [
    # ★ emoji and triggers
    (r'★\s*', ''),
    (r'喵喵喵', ''),
    # AI/internal markers
    (r'Mavis\s+', ''),
    (r'老刀[^\s]*\s*', ''),
    # [E3 Visium] / [E6 MR] / [E8 PDO] / [v6.7 NEW] / [v11 NEW] tags in captions
    (r'\[E3 Visium\]\s*', ''),
    (r'\[E6 MR\]\s*', ''),
    (r'\[E8 PDO\]\s*', ''),
    (r'\[v\d+\.?\d*\s*NEW\]\s*', ''),
    (r'\[v11\]\s*', ''),
    # Version markers in subtitle / context
    (r'Manuscript\s+v\d+\.?\d*\s*[—–-]\s*', 'Manuscript — '),
    (r'\(v\d+\.?\d*\s*NEW\)', ''),
    # v11/v8/v6.7 bare references in narrative
    (r'\bv11\s+final\b', 'the present study'),
    (r'\bv11\b(?!\.)', 'the present study'),
    (r'\bv8\b(?!\.)', 'the present study'),
    (r'\bv6\.7\b(?!\.)', 'the present study'),
    # Star markers before important sentences
    (r'^\s*\*\s*', ''),
    # "trigger" / "click" / "hot-fix" 词
    (r'\btrigger\b', 'feedback'),
    (r'\b拍板\b', 'decision'),
    (r'\b老刀\s*trigger', 'revision'),
]

# Apply to all paragraphs
for p in doc.paragraphs:
    txt = p.text
    new_txt = txt
    for pat, repl in cleanup_patterns:
        new_txt = re.sub(pat, repl, new_txt)
    if new_txt != txt:
        p.text = new_txt

print(f'[v8 polish] text cleanup applied to all paragraphs')

# ============ STEP 2: MOVE FIGURES INLINE ============
# Find figure paragraphs (img=True) and their caption paragraphs
# Then map figures to subsections (§4.1-4.5) and move them inline

paragraphs = list(doc.paragraphs)

def has_image(p):
    return any(r.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip') for r in p.runs)

# Categorize each paragraph
img_paragraphs = []  # list of (idx, element, caption_idx_or_None)
last_fig_caption_idx = None
for i, p in enumerate(paragraphs):
    if has_image(p):
        img_paragraphs.append((i, p._p, None))
    elif p.text.strip().startswith('Figure ') and re.match(r'^Figure \d+', p.text.strip()):
        # This is a figure caption
        if img_paragraphs and img_paragraphs[-1][2] is None:
            img_paragraphs[-1] = (img_paragraphs[-1][0], img_paragraphs[-1][1], i)

print(f'[v8 polish] found {len(img_paragraphs)} figure (image + caption) pairs')

# Map figure number → subsection
def subsection_for_figure(caption_text):
    """Return which §4.x subsection this figure belongs to."""
    m = re.match(r'Figure (\d+)', caption_text)
    if not m:
        return None
    n = int(m.group(1))
    if 1 <= n <= 6:
        return '4.1'
    elif 7 <= n <= 22:
        return '4.2'
    elif 23 <= n <= 30:
        return '4.3'
    elif 31 <= n <= 40:
        return '4.4'
    elif 41 <= n <= 46:
        return '4.5'
    return None

# Find the section anchor paragraphs (where each §4.x ends)
# §4.x subsections are inside §4 Results block. After each §4.x body paragraph, we want to insert the relevant figures.
# The end of §4.5 body is where we currently have all figures stacked. Strategy: move figures from this block to after their respective subsection's body paragraph.

# Find the body paragraph for each §4.x subsection
subsection_body_idx = {}  # '4.1' → idx of body paragraph
for i, p in enumerate(paragraphs):
    if p.style.name == 'Heading 2' and p.text.strip().startswith('4.'):
        sub_num = p.text.strip().split()[0]  # '4.1' etc
        # Find the next non-heading paragraph as body
        for j in range(i + 1, len(paragraphs)):
            if paragraphs[j].style.name == 'Heading 2' or (paragraphs[j].style.name == 'Heading 1' and paragraphs[j].text.strip().startswith('5.')):
                break
            if paragraphs[j].text.strip():
                subsection_body_idx[sub_num] = j
                break

print(f'[v8 polish] subsection body indices: {subsection_body_idx}')

# Now, for each figure, find its destination
figure_moves = []  # list of (figure_xml_element, caption_paragraph_idx, target_subsection)
for img_idx, img_elem, cap_idx in img_paragraphs:
    if cap_idx is None:
        continue
    cap_text = paragraphs[cap_idx].text.strip()
    sub = subsection_for_figure(cap_text)
    if sub is None:
        continue
    figure_moves.append((img_idx, img_elem, cap_idx, paragraphs[cap_idx]._p, sub))

print(f'[v8 polish] {len(figure_moves)} figures to move')

# To avoid index shifting during moves, work with XML elements directly.
# Strategy: remove figure+caption from current location, insert after target subsection body.

# First, remove all figure (img) + caption pairs from their current block
for img_idx, img_elem, cap_idx, cap_elem, sub in figure_moves:
    img_elem.getparent().remove(img_elem)
    cap_elem.getparent().remove(cap_elem)

print(f'[v8 polish] removed {len(figure_moves)} figures from old location')

# Now re-find subsection body positions (indices changed)
paragraphs = list(doc.paragraphs)
subsection_body_idx = {}
for i, p in enumerate(paragraphs):
    if p.style.name == 'Heading 2' and p.text.strip().startswith('4.'):
        sub_num = p.text.strip().split()[0]
        for j in range(i + 1, len(paragraphs)):
            if paragraphs[j].style.name == 'Heading 2' or (paragraphs[j].style.name == 'Heading 1' and paragraphs[j].text.strip().startswith('5.')):
                break
            if paragraphs[j].text.strip():
                subsection_body_idx[sub_num] = j
                break

# Group figure moves by subsection
from collections import defaultdict
moves_by_sub = defaultdict(list)
for img_idx, img_elem, cap_idx, cap_elem, sub in figure_moves:
    moves_by_sub[sub].append((img_elem, cap_elem))

# For each subsection, find the body paragraph and insert figures AFTER it
for sub, moves in moves_by_sub.items():
    if sub not in subsection_body_idx:
        print(f'  WARN: subsection {sub} not found')
        continue
    body_p = paragraphs[subsection_body_idx[sub]]._p
    # Insert each figure (img then caption) AFTER body_p
    for img_elem, cap_elem in moves:
        # insert img after body_p
        body_p.addnext(img_elem)
        # insert caption after img
        img_elem.addnext(cap_elem)
        # update body_p reference for next iteration
        body_p = cap_elem

print(f'[v8 polish] figures moved inline to subsections')

# ============ STEP 3: SAVE ============
print(f'[v8 polish] saving to {V8}')
doc.save(V8)

# Verify
import os
print(f'[v8 polish] v8 size: {os.path.getsize(V8)/1024/1024:.1f} MB')

# Word count
doc2 = Document(V8)
all_text = '\n'.join(p.text for p in doc2.paragraphs if p.text.strip())
print(f'[v8 polish] paragraphs: {len(doc2.paragraphs)}, words: {len(all_text.split())}')

# Check for residual AI markers
residuals = []
for p in doc2.paragraphs:
    t = p.text
    if '喵喵喵' in t or 'Mavis' in t or '老刀' in t:
        residuals.append(t[:80])
print(f'[v8 polish] residual AI markers: {len(residuals)}')
for r in residuals[:5]:
    print(f'  → {r}')

# Check version markers
ver_residuals = []
for p in doc2.paragraphs:
    t = p.text
    if re.search(r'\bv\d+\.?\d*\b', t):
        ver_residuals.append(t[:80])
print(f'[v8 polish] residual version markers: {len(ver_residuals)}')
for r in ver_residuals[:5]:
    print(f'  → {r}')