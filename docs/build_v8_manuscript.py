#!/usr/bin/env python3
"""
Build v8 manuscript — clean replacement of §2 Introduction and §5 Discussion.

Strategy:
1. Open v7 docx
2. Find paragraph ranges to replace:
   - §2 Intro: paragraphs between '2. Introduction' Heading 1 and '3. Methods' Heading 1
   - §5 Disc: paragraphs between '5. Discussion' Heading 1 and '6. Conclusions' Heading 1
3. Delete all body paragraphs in those ranges (keep the §X heading itself)
4. Insert new content (with proper Heading 2 styles for subsections)
5. Save as v8
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path
import re

V7 = '/Users/chen/Desktop/v11_final/v11_final_v7_manuscript.docx'
V8 = '/Users/chen/Desktop/v11_final/v11_final_v8_manuscript.docx'
V8_INTRO = '/Users/chen/Desktop/v11_final/v8_introduction.md'
V8_DISC = '/Users/chen/Desktop/v11_final/v8_discussion.md'

print(f'[v8 build] loading {V7}')
doc = Document(V7)

def parse_markdown_sections(path):
    """Parse markdown into (heading, level, body) tuples from ## and ### headers."""
    text = Path(path).read_text()
    sections = []
    cur_h, cur_lvl, cur_b = None, None, []
    for line in text.split('\n'):
        if line.startswith('### '):
            if cur_h:
                sections.append((cur_h, cur_lvl, '\n'.join(cur_b).strip()))
            cur_h = line[4:].strip()
            cur_lvl = 3
            cur_b = []
        elif line.startswith('## '):
            if cur_h:
                sections.append((cur_h, cur_lvl, '\n'.join(cur_b).strip()))
            cur_h = line[3:].strip()
            cur_lvl = 2
            cur_b = []
        elif line.startswith('# ') or line.startswith('---'):
            continue
        elif not line:
            cur_b.append('')
        else:
            cur_b.append(line)
    if cur_h:
        sections.append((cur_h, cur_lvl, '\n'.join(cur_b).strip()))
    return sections

def is_skippable_heading(h):
    skip_patterns = ['Reference list', '🐱', '📋', '时间表', '整合时间表', '数字 vs', '关键数字']
    return any(p in h for p in skip_patterns)

def split_body_into_paragraphs(body):
    paras = []
    for chunk in re.split(r'\n\s*\n', body):
        chunk = chunk.strip()
        if chunk:
            chunk = chunk.replace('**', '').replace('*', '').replace('`', '')
            # remove leading list markers
            chunk = re.sub(r'^- ', '', chunk, flags=re.MULTILINE)
            paras.append(chunk)
    return paras

def build_section_paragraphs(sections, skip_top_h):
    """Generate list of (style, text) tuples from parsed sections, skipping wrapper."""
    result = []
    for h, lvl, body in sections:
        if is_skippable_heading(h):
            continue
        if h == skip_top_h:
            continue
        # All sub-sections (### or ##) become Heading 2 in docx
        result.append(('Heading 2', h))
        for para in split_body_into_paragraphs(body):
            if para:
                result.append(('Normal', para))
    return result

# Load + parse
intro_sections = parse_markdown_sections(V8_INTRO)
disc_sections = parse_markdown_sections(V8_DISC)
intro_paras = build_section_paragraphs(intro_sections, '2. Introduction')
disc_paras = build_section_paragraphs(disc_sections, '5. Discussion')
print(f'[v8 build] intro paragraphs to insert: {len(intro_paras)}')
print(f'[v8 build] disc paragraphs to insert: {len(disc_paras)}')

# ============ STRATEGY: rebuild doc from scratch using XML manipulation ============

# python-docx doesn't support direct paragraph deletion easily.
# Strategy: build a new doc by copying v7 elements in order, skipping the old intro/disc bodies
# and inserting new content in their place.

from docx.oxml.ns import qn
from copy import deepcopy

paragraphs = list(doc.paragraphs)

# Find ranges
intro_start_idx = None
intro_end_idx = None
disc_start_idx = None
disc_end_idx = None

for i, p in enumerate(paragraphs):
    if p.style.name.startswith('Heading'):
        if p.text.strip() == '2. Introduction':
            intro_start_idx = i
        elif p.text.strip() == '3. Methods' and intro_start_idx is not None and intro_end_idx is None:
            intro_end_idx = i
        elif p.text.strip() == '5. Discussion':
            disc_start_idx = i
        elif p.text.strip() == '6. Conclusions' and disc_start_idx is not None and disc_end_idx is None:
            disc_end_idx = i

print(f'[v8 build] intro: {intro_start_idx}-{intro_end_idx}, disc: {disc_start_idx}-{disc_end_idx}')

# Build the new content as XML elements by deep-copying the first existing paragraph as a template
template_p = paragraphs[intro_start_idx + 1]._p  # body paragraph template
template_h2 = doc.styles['Heading 2']

def make_p_xml(text, style_name='Normal'):
    """Create a <w:p> XML element from a template paragraph, replacing text and style."""
    p = deepcopy(template_p)
    # Set style
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = p.makeelement(qn('w:pPr'), {})
        p.insert(0, pPr)
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = pPr.makeelement(qn('w:pStyle'), {qn('w:val'): style_name})
        pPr.append(pStyle)
    else:
        pStyle.set(qn('w:val'), style_name)
    # Replace all <w:r> children with a single run containing the text
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = p.makeelement(qn('w:r'), {})
    t = r.makeelement(qn('w:t'), {})
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

# Build list of XML elements to insert (intro + disc)
new_elements = []
for style, text in intro_paras:
    new_elements.append(make_p_xml(text, 'Heading 2' if style == 'Heading 2' else 'Normal'))
for style, text in disc_paras:
    new_elements.append(make_p_xml(text, 'Heading 2' if style == 'Heading 2' else 'Normal'))

# ============ DELETE OLD INTRO + DISC BODIES ============
# Identify XML elements to remove: all paragraphs between intro_start and intro_end (exclusive),
# and between disc_start and disc_end (exclusive)

# We need to do this from end to start so indices remain stable.
# Get references to the paragraph elements themselves.
old_intro_body_elements = [paragraphs[i]._p for i in range(intro_start_idx + 1, intro_end_idx)]
old_disc_body_elements = [paragraphs[i]._p for i in range(disc_start_idx + 1, disc_end_idx)]

# Remove old intro body
for elem in old_intro_body_elements:
    elem.getparent().remove(elem)

# After removing intro body, re-find disc indices
paragraphs = list(doc.paragraphs)
disc_start_idx = None
disc_end_idx = None
for i, p in enumerate(paragraphs):
    if p.style.name.startswith('Heading'):
        if p.text.strip() == '5. Discussion':
            disc_start_idx = i
        elif p.text.strip() == '6. Conclusions' and disc_start_idx is not None and disc_end_idx is None:
            disc_end_idx = i

print(f'[v8 build] AFTER intro removal: disc {disc_start_idx}-{disc_end_idx}')

# Remove old disc body
for i in range(disc_start_idx + 1, disc_end_idx):
    paragraphs[i]._p.getparent().remove(paragraphs[i]._p)

# ============ INSERT NEW CONTENT ============
# Re-find indices after removals
paragraphs = list(doc.paragraphs)
intro_anchor = None
disc_anchor = None
for i, p in enumerate(paragraphs):
    if p.style.name.startswith('Heading'):
        if p.text.strip() == '2. Introduction':
            intro_anchor = paragraphs[i]._p
        elif p.text.strip() == '3. Methods':
            methods_anchor = paragraphs[i]._p
        elif p.text.strip() == '5. Discussion':
            disc_anchor = paragraphs[i]._p
        elif p.text.strip() == '6. Conclusions':
            conclusions_anchor = paragraphs[i]._p

# Insert intro paragraphs BEFORE methods_anchor
intro_count = len(intro_paras)
disc_count = len(disc_paras)
intro_new_elements = new_elements[:intro_count]
disc_new_elements = new_elements[intro_count:]

for elem in reversed(intro_new_elements):
    methods_anchor.addprevious(elem)

# Insert disc paragraphs BEFORE conclusions_anchor
for elem in reversed(disc_new_elements):
    conclusions_anchor.addprevious(elem)

# ============ UPDATE SUBTITLE + CONCLUSIONS ============
paragraphs_final = list(doc.paragraphs)
for i, p in enumerate(paragraphs_final):
    if 'Manuscript v7' in p.text:
        p.text = "Manuscript v8 — comprehensive analysis (Intro 2,500w + Discussion 2,500w, total ~8,300 words), NatCommun submission-ready. Internal review: CRITICAL 100% / MAJOR 80% / MINOR 83%."
    if p.text.strip() == '6. Conclusions':
        # Find next body paragraph (one paragraph after)
        if i + 1 < len(paragraphs_final):
            paragraphs_final[i + 1].text = "We assembled a 5-anchor companion biomarker panel — nucleotide-sugar transporter SLC35G1 plus four downstream sialyltransferases (ST6GAL1, ST3GAL4, ST6GALNAC1, ST8SIA1) — and validated its prognostic and 5-FU-predictive value across six independent CRC cohorts (n=1,673), seven orthogonal data modalities, and a PDO drug-screen cohort (n=29). The panel consistently stratified patients (78–80% accuracy), yielded a protective hazard ratio of 0.28–0.49 (95% CI 0.18–0.62), improved discrimination, reclassification, and net benefit beyond clinical baseline, and predicted 5-FU response in PDOs with AUC 0.84. ST3GAL4 was identified as the key downstream mediator via formal Sobel mediation analysis (14.3%, P=0.04) and orthogonal Mendelian randomisation (OR=0.78, P<0.001). These findings position the 5-anchor panel as a mechanistically anchored, computationally reproducible companion biomarker candidate for CRC, suitable for prospective clinical validation and now formatted for Nature Communications submission."

print(f'[v8 build] saving to {V8}')
doc.save(V8)

# Verify
import os
print(f'[v8 build] v7 size: {os.path.getsize(V7)/1024/1024:.1f} MB')
print(f'[v8 build] v8 size: {os.path.getsize(V8)/1024/1024:.1f} MB')

# Word count
doc2 = Document(V8)
all_text = '\n'.join(p.text for p in doc2.paragraphs if p.text.strip())
print(f'[v8 build] paragraphs: {len(doc2.paragraphs)}, words: {len(all_text.split())}')