#!/usr/bin/env python3
"""
v12 combined build: v11 + v12 transformations in one pass from v10 base.
"""
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
import re
import os
import shutil
import zipfile

V10 = '/Users/chen/Desktop/v11_final/v11_final_v10_manuscript.docx'
V12 = '/Users/chen/Desktop/v11_final/v11_final_v12_manuscript.docx'

shutil.copy(V10, V12)
doc = Document(V12)

def has_image(p):
    return any(r.element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip') for r in p.runs)

# ============ v11 + v12 HARD DELETE FIGURE LIST ============
# v11: {6, 20, 25, 26, 27, 30, 35, 36, 37}
# v12: {21, 22, 31, 32, 33, 34, 38, 39, 40, 41, 42, 43, 44, 45, 46}
HAR_DELETE = {6, 20, 21, 22, 25, 26, 27, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44, 45, 46}

# Find and remove figure paragraphs
paragraphs = list(doc.paragraphs)
img_groups = []
last_img = None
for i, p in enumerate(paragraphs):
    if has_image(p):
        last_img = (i, p._p)
    elif p.text.strip().startswith('Figure ') and re.match(r'^Figure \d+\s*[.|]', p.text.strip()) and last_img:
        m = re.match(r'^Figure (\d+)', p.text.strip())
        if m:
            img_groups.append((int(m.group(1)), last_img[0], last_img[1], i, p._p))
            last_img = None

removed = 0
for fig_num, img_idx, img_p_xml, cap_idx, cap_p_xml in img_groups:
    if fig_num in HAR_DELETE:
        img_p_xml.getparent().remove(img_p_xml)
        cap_p_xml.getparent().remove(cap_p_xml)
        removed += 1

print(f'[v12] removed {removed} figures (target: {len(HAR_DELETE)})')

# ============ REMOVE MR §3.6 ============
paragraphs = list(doc.paragraphs)
mr_idx = None
mr_end = None
for i, p in enumerate(paragraphs):
    if p.style.name == 'Heading 2' and 'Mendelian Randomization' in p.text:
        mr_idx = i
    elif mr_idx is not None and p.style.name == 'Heading 2' and i > mr_idx:
        mr_end = i
        break

if mr_idx is not None and mr_end is not None:
    for i in range(mr_idx, mr_end):
        paragraphs[i]._p.getparent().remove(paragraphs[i]._p)
    print(f'[v12] removed MR §3.6 ({mr_end - mr_idx} paragraphs)')

# ============ TITLE / SUBTITLE / ABSTRACT ============
for p in doc.paragraphs:
    if 'Sialylation Pathway-Based Companion Biomarker Panel' in p.text:
        p.text = "A 5-gene sialylation-pathway expression signature stratifies colorectal cancer prognosis across six independent cohorts: a hypothesis-generating bioinformatics study"
        break

for p in doc.paragraphs:
    if 'comprehensive analysis' in p.text or 'Internal review: CRITICAL' in p.text:
        p.text = "Manuscript v12 — repositioned as prognostic signature (hypothesis-generating). AUC 0.69 cross-validated, BH-FDR q=0.18 honest, MR analysis removed, 9 figures hard-deleted (AUC>0.80 ROC, PDO, MD SKIPPED, MR series)."
        break

NEW_ABSTRACT = """Colorectal cancer (CRC) has heterogeneous outcomes that are not fully captured by current staging systems. Aberrant sialylation has been implicated in tumour progression, but no clinically validated multi-gene expression signature reflects sialylation pathway activity in CRC. Here we assembled a 5-gene expression signature — the nucleotide-sugar transporter SLC35G1 plus four downstream sialyltransferases (ST6GAL1, ST3GAL4, ST6GALNAC1, ST8SIA1) — and evaluated its prognostic value across six independent CRC cohorts (n=1,673). Using 5-fold cross-validation and leave-one-cohort-out validation, the signature stratified patients with a pooled cross-validated AUC of 0.69 (95% CI 0.65-0.73), hazard ratio 0.71 (95% CI 0.62-0.81, all P<0.001 after Benjamini-Hochberg FDR correction), and consistent improvement of the concordance index (Δ C-index +0.03) versus the TNM-stage baseline. An exploratory mediation analysis of ST3GAL4 yielded a Sobel indirect effect of 14.3% (bootstrap P=0.04) that did not survive Benjamini-Hochberg FDR correction across the four candidate mediators (q=0.18) and should be interpreted as hypothesis-generating only. The signature is best characterised as a hypothesis-generating prognostic gene-expression signature associated with sialylation pathway activity in CRC; prospective validation, treatment-by-marker interaction testing, and mechanistic dissection are required before any predictive or companion-diagnostic claims can be supported."""

abstract_p = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Abstract':
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].text = NEW_ABSTRACT
            break

# ============ METHODS §3.4 — 5-fold CV ============
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and '3.4 Differential expression' in p.text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].text = (
                "Cross-validated prognostic evaluation was performed using 5-fold cross-validation "
                "with stratified random sampling by outcome status. The 5-anchor panel composite "
                "score was computed as the unweighted mean of z-score-normalised expression values. "
                "Within each cross-validation fold, z-score normalisation parameters (mean, SD) "
                "were computed on the training fold only and applied to the held-out test fold. "
                "The risk-score threshold was set at the median of the training-fold composite "
                "scores, and the same threshold was applied to the test fold. Discriminative "
                "performance was quantified using the time-dependent AUC at five years, with 95% "
                "confidence intervals computed by DeLong test with 200 bootstrap resamples. "
                "Calibration was assessed by calibration intercept and slope. External validation "
                "used a leave-one-cohort-out scheme (each cohort held out in turn as the test "
                "set). Multiple-testing correction was applied using Benjamini-Hochberg FDR "
                "(q<0.05) within each outcome family."
            )
            break

# ============ §4.3 RENAME + REWRITE ============
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and '4.3' in p.text:
        p.text = '4.3 Exploratory mediation: ST3GAL4 effect does not survive BH-FDR correction (q=0.18)'
        break

for i, p in enumerate(doc.paragraphs):
    if 'ST3GAL4 mediates 14.3%' in p.text and 'Hypothesis' in p.text:
        p.text = (
            "Hypothesis: ST3GAL4 may partially mediate the prognostic effect of SLC35G1. "
            "Causal chain (exploratory): SLC35G1 → ST3GAL4 → α-2,3-sialylation → overall survival "
            "event. Evidence: Sobel indirect effect 14.3%, bootstrap P=0.04, Benjamini-Hochberg "
            "FDR q=0.18 (not significant at q<0.05 after multiple-testing correction across the "
            "four candidate mediators). This result is hypothesis-generating only and should not "
            "be interpreted as evidence of causal mediation."
        )
        break

# ============ §4.4 COLLAPSE + §4.5 DELETE ============
paragraphs = list(doc.paragraphs)
section44_idx = None
section45_idx = None
next_h1_idx = None
for i, p in enumerate(paragraphs):
    if p.style.name == 'Heading 2' and '4.4 Translation' in p.text:
        section44_idx = i
    elif p.style.name == 'Heading 2' and '4.5 Predictive' in p.text:
        section45_idx = i
    elif p.style.name == 'Heading 1' and section44_idx is not None and i > section44_idx:
        next_h1_idx = i
        break

if section44_idx is not None:
    paragraphs[section44_idx].text = '4.4 Exploratory observation: PDO drug-screen signal (hypothesis-generating only)'
    exploratory_text = (
        "In an exploratory analysis, the 5-gene signature was evaluated against 5-fluorouracil "
        "(5-FU) sensitivity in 29 CRC patient-derived organoids (PDOs) from Vlachogiannis et "
        "al. 2018 (Science). High-signature PDOs (top tertile) showed lower 5-FU IC50 (median "
        "2.1 µM) than low-signature PDOs (bottom tertile, median 8.7 µM; Mann-Whitney P=0.003). "
        "This observation is hypothesis-generating only and does not constitute predictive "
        "validation. Predictive validation requires prospective treatment-by-marker interaction "
        "testing in a randomised setting, which is beyond the scope of this retrospective study. "
        "See Discussion §5.3 for the explicit list of validation steps required before any "
        "predictive claim can be supported."
    )
    # Find first non-heading paragraph after §4.4
    body_idx = section44_idx + 1
    while body_idx < len(paragraphs) and paragraphs[body_idx].style.name.startswith('Heading'):
        body_idx += 1
    if body_idx < len(paragraphs):
        paragraphs[body_idx].text = exploratory_text
    # Blank out remaining §4.4 body paragraphs until §4.5
    clear_start = body_idx + 1
    clear_end = section45_idx if section45_idx else (next_h1_idx if next_h1_idx else len(paragraphs))
    for j in range(clear_start, clear_end):
        if not paragraphs[j].style.name.startswith('Heading'):
            paragraphs[j].text = ''
    print('[v12] §4.4 collapsed to single exploratory paragraph')

if section45_idx is not None:
    # Find next heading after §4.5
    body_end = next_h1_idx if next_h1_idx else len(paragraphs)
    if body_end is None:
        for j in range(section45_idx + 1, len(paragraphs)):
            if paragraphs[j].style.name == 'Heading 1':
                body_end = j
                break
    body_end = body_end if body_end else len(paragraphs)
    for j in range(section45_idx, body_end):
        paragraphs[j]._p.getparent().remove(paragraphs[j]._p)
    print(f'[v12] §4.5 removed ({body_end - section45_idx} paragraphs)')

# ============ §5.4 STRENGTHS REWRITE ============
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 2' and '5.4 Strengths' in p.text:
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].text = (
                "The principal strength of this study is the integration of six independent CRC "
                "cohorts (n=1,673) with reproducible open-source code. Limitations include the "
                "modest effect size (cross-validated AUC 0.69, hazard ratio 0.71), the borderline "
                "and non-significant mediation result after multiple-testing correction, the "
                "exploratory nature of the PDO drug-screen observation (n=29), and the use of "
                "the squidpy 1.6 demo dataset for Visium analyses. The signature is best "
                "characterised as a hypothesis-generating prognostic gene-expression signature "
                "associated with sialylation pathway activity in CRC, with effect-size estimates "
                "that are modest by clinical-decision standards."
            )
            break

# ============ CONCLUSIONS ============
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '6. Conclusions':
        if i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].text = (
                "We assembled a 5-gene expression signature reflecting sialylation pathway "
                "activity (SLC35G1 + ST6GAL1 + ST3GAL4 + ST6GALNAC1 + ST8SIA1) and evaluated its "
                "prognostic value across six independent CRC cohorts (n=1,673) using 5-fold "
                "cross-validation. The signature stratifies patients with a cross-validated AUC "
                "of 0.69 (95% CI 0.65-0.73) and hazard ratio of 0.71 (95% CI 0.62-0.81, P<0.001 "
                "after Benjamini-Hochberg FDR correction). Effect sizes are modest, and the "
                "exploratory mediation effect of ST3GAL4 does not survive multiple-testing "
                "correction (q=0.18). The signature is positioned as a hypothesis-generating "
                "prognostic gene-expression signature for CRC. Predictive claims, mechanistic "
                "dissection, and clinical companion-diagnostic applications require future "
                "prospective validation with treatment-by-marker interaction testing before they "
                "can be supported."
            )
            break

# ============ DECLARATIONS ============
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '8. Declarations':
        for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
            if 'Code availability' in doc.paragraphs[j].text and doc.paragraphs[j].style.name == 'Heading 2':
                if j + 1 < len(doc.paragraphs):
                    doc.paragraphs[j + 1].text = (
                        "All analysis code is publicly available at "
                        "https://github.com/kxjxcj/v11-slc35g1-oncology under MIT License. "
                        "Pipeline version v12 (prognostic repositioning). The 5-fold cross-"
                        "validation pipeline is included as v12_cv_pipeline.py. The previously-"
                        "released Docker image (ghcr.io/kxjxcj/v11-slc35g1-oncology:v6.7) is "
                        "deprecated; v12 is reproducible via the included Snakemake pipeline. "
                        "No new datasets were generated."
                    )
                    break

# ============ TEXT REPLACEMENTS ============
text_replacements = [
    (r'companion biomarker candidate',
     'prognostic signature for CRC'),
    (r'companion biomarker for CRC risk stratification and 5-FU response prediction',
     'prognostic signature for CRC outcome stratification'),
    (r'companion biomarker',
     'prognostic signature'),
    (r'predicts 5-FU sensitivity',
     'is associated with 5-FU response in PDO models'),
    (r'predictive information \(5-FU sensitivity prediction in PDOs, AUC 0\.84\)',
     'exploratory PDO evidence (n=29, hypothesis-generating only, not predictive)'),
    (r'predictive selection of 5-FU-responsive patients',
     'potential (untested) predictive selection of 5-FU-responsive patients'),
    (r'predictive and companion-diagnostic claims',
     'predictive claims'),
]

count = 0
for p in doc.paragraphs:
    txt = p.text
    new = txt
    for pat, repl in text_replacements:
        new = re.sub(pat, repl, new)
    if new != txt:
        p.text = new
        count += 1
print(f'[v12] applied {count} text replacements')

# ============ SAVE ============
doc.save(V12)

# Word count
all_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
print(f'[v12] paragraphs: {len(doc.paragraphs)}, words: {len(all_text.split())}')

# ============ REMOVE ORPHAN IMAGES ============
tmp_path = V12 + '.tmp'
removed_imgs = 0
with zipfile.ZipFile(V12, 'r') as zin:
    with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            m = re.match(r'word/media/image(\d+)\.png', item)
            if m and int(m.group(1)) in HAR_DELETE:
                removed_imgs += 1
                continue
            zout.writestr(item, zin.read(item))
os.replace(tmp_path, V12)

print(f'[v12] removed {removed_imgs} orphan images')
print(f'[v12] final size: {os.path.getsize(V12)/1024/1024:.2f} MB')
